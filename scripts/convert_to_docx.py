"""将 delivery-implementation-report.md 转换为 docx 格式"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

def create_docx():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题样式设置
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = '微软雅黑'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        heading_font.bold = True
        if i == 1:
            heading_font.size = Pt(22)
            heading_font.color.rgb = RGBColor(0, 51, 102)
        elif i == 2:
            heading_font.size = Pt(16)
            heading_font.color.rgb = RGBColor(0, 76, 153)
        elif i == 3:
            heading_font.size = Pt(14)
            heading_font.color.rgb = RGBColor(0, 102, 153)

    # 封面
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Coding Agent 交付实现结果报告')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Capybara 项目')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(102, 102, 102)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_paragraph()
    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run('2026年6月')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(102, 102, 102)
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    doc.add_page_break()

    # 目录页
    doc.add_heading('目录', level=1)
    toc_items = [
        '1. 项目概述',
        '2. 核心架构',
        '3. 已实现功能清单',
        '   3.1 核心基础设施（core/）',
        '   3.2 工具层（tools/）',
        '   3.3 Agent 实现（agents/）',
        '4. 测试覆盖情况',
        '5. Demo 示例',
        '6. 验证记录',
        '7. 项目亮点',
        '8. 后续规划',
        '9. 总结',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)

    doc.add_page_break()

    # 第1章：项目概述
    doc.add_heading('1. 项目概述', level=1)
    doc.add_paragraph(
        '本项目是一个从零搭建的 Python 编程 Agent 核心库，能够读代码、改代码、跑验证，'
        '并具备完整的生产化能力。项目采用分层架构设计，代码分布在四个顶层包：'
        'core/（基础设施）、tools/（工具层）、agents/（Agent 实现）、tests/（测试）。'
    )

    doc.add_heading('技术栈', level=2)
    table = doc.add_table(rows=4, cols=2, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = table.rows[0].cells
    cells[0].text = '项目'
    cells[1].text = '说明'
    cells = table.rows[1].cells
    cells[0].text = '语言'
    cells[1].text = 'Python 3.12'
    cells = table.rows[2].cells
    cells[0].text = '运行环境'
    cells[1].text = 'conda agent 环境'
    cells = table.rows[3].cells
    cells[0].text = '外部依赖'
    cells[1].text = 'openai、python-dotenv、pydantic'

    doc.add_paragraph()

    # 第2章：核心架构
    doc.add_heading('2. 核心架构', level=1)
    doc.add_paragraph(
        '项目采用分层架构设计，CodingAgent 位于最上层，调用 core/ 和 tools/ 提供的能力。'
    )

    # 架构图
    doc.add_heading('架构图', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '┌─────────────────────────────────────────────────────────────┐\n'
        '│                     CodingAgent (agents/)                    │\n'
        '│  ReAct 循环 │ 多轮记忆 │ 验证闭环 │ 子 Agent │ MCP 适配      │\n'
        '└─────────────────────────────────────────────────────────────┘\n'
        '                              │\n'
        '        ┌─────────────────────┼─────────────────────┐\n'
        '        ▼                     ▼                     ▼\n'
        '┌──────────────┐     ┌──────────────┐     ┌──────────────┐\n'
        '│  core/       │     │  tools/      │     │  tests/      │\n'
        '│  Agent 基类  │◄────│  工具抽象    │     │  246 个用例  │\n'
        '│  LLM 接入    │     │  14 个工具   │     │  100% 通过   │\n'
        '│  会话/Trace  │     │  代码智能    │     │  零外部依赖  │\n'
        '└──────────────┘     └──────────────┘     └──────────────┘'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading('包职责划分', level=2)
    table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['包', '职责', '关键模块']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    data = [
        ['core/', '基础设施层：Agent 基类、LLM 接入、消息、配置、会话持久化、流式事件、trace、上下文预算、稳健性、MCP 客户端', 'agent.py、llm.py、config.py、trace.py、resilience.py、mcp.py'],
        ['tools/', '工具层：Tool 抽象 + 注册表 + 14 个编程工具（含 code_intel/ 代码智能子包）', 'base.py、registry.py、code_intel/'],
        ['agents/', '具体 Agent 实现：CodingAgent + 子 Agent / 项目检测 / 改动追踪 / 待办 / 开发日志 / 技能 / MCP 工具包装', 'coding_agent.py、subagent.py、verify.py'],
        ['tests/', 'stdlib unittest 单元测试（无需 pytest），246 个用例，全部使用本地 StubLLM，不发起真实 API 请求', '12 个测试模块'],
    ]
    for i, row_data in enumerate(data, 1):
        for j, cell_data in enumerate(row_data):
            table.rows[i].cells[j].text = cell_data

    doc.add_page_break()

    # 第3章：已实现功能清单
    doc.add_heading('3. 已实现功能清单', level=1)

    # 3.1 核心基础设施
    doc.add_heading('3.1 核心基础设施（core/）', level=2)

    core_modules = [
        ('Agent 基类 (core/agent.py)', [
            '抽象基类设计，持有 name / llm / system_prompt / config / tool_registry',
            '消息历史管理（list[Message]）',
            '同步执行入口 run()',
            '异步执行入口 arun()（线程池包装）',
            '异步流式入口 arun_stream()（粗粒度生命周期事件）',
            '会话持久化（save_session / load_session / list_sessions）',
        ]),
        ('LLM 接入 (core/llm.py + core/llm_adapters.py)', [
            'AgentsLLM 门面：从 .env 读取 LLM_MODEL_ID / LLM_API_KEY / LLM_BASE_URL',
            '统一接口：invoke / stream_invoke / invoke_with_tools 及异步包装',
            'OpenAI 兼容适配器（支持 DeepSeek / Qwen / Kimi 等）',
            'thinking model 的 reasoning_content 支持',
            '原生异步客户端 astream_invoke',
        ]),
        ('消息与配置 (core/message.py + core/config.py)', [
            'Pydantic 消息模型（user / assistant / system / tool / summary）',
            '消息序列化（to_dict / from_dict / to_text）',
            '集中配置模型，涵盖 LLM、上下文压缩、工具截断、trace、skills、熔断器、会话持久化、子代理等开关',
        ]),
        ('会话持久化 (core/session_store.py)', [
            'JSON 原子持久化（temp + os.replace()）',
            'save / load / list_sessions / delete',
            '配置和工具 schema 一致性检查',
            'delete() 入参容错（裸名/.json/完整路径均可）',
        ]),
        ('流式事件 (core/streaming.py + core/lifecycle.py)', [
            'StreamEvent，to_sse() / to_dict()',
            'StreamBuffer 带背压控制',
            'EventType：Agent/Step/LLM/Tool/Thinking 等 16 种事件类型',
            'AgentEvent、ExecutionContext、LifecycleHook',
        ]),
        ('可观测性 Trace (core/trace.py)', [
            '轻量 JSONL 执行追踪',
            'start_run() 开启一次运行',
            'record(event, **data) 逐事件追加写入',
            'write_html_report() 导出 HTML 报告',
            '敏感信息脱敏（key 名精确匹配 + sk-/Bearer 正则）',
            '超长字段截断',
            '所有写入异常被吞，作为旁路观测绝不影响主流程',
        ]),
        ('上下文预算管理 (core/context_manager.py)', [
            'estimate_tokens(obj)：token 粗估（len//4）',
            'truncate_output(text, max_lines, max_bytes, direction)：截断单次输出',
            'compact_history(messages, ...)：历史压缩',
            '只动无 tool 的跨轮历史，不破坏 assistant↔tool 配对',
        ]),
        ('稳健性 (core/resilience.py)', [
            'CircuitBreaker：三态熔断器（closed/open/half_open）',
            'retry_call(fn, *, max_retries, backoff, on_retry)：指数退避重试',
            'aretry_call(make_coro, ...)：异步重试',
            '熔断异常不重试',
        ]),
        ('MCP 客户端 (core/mcp.py)', [
            '零依赖 MCP（Model Context Protocol）stdio 客户端',
            'start() 启动 server 子进程并完成 initialize 握手',
            'list_tools() / call_tool(name, args)',
            'close() 终止进程，支持 with 上下文',
            '行分隔 JSON-RPC 2.0，后台读线程 + 队列匹配响应',
            '仅用标准库（subprocess/json/threading/queue），零新增依赖',
        ]),
    ]

    for module_name, features in core_modules:
        doc.add_heading(module_name, level=3)
        for feature in features:
            p = doc.add_paragraph(feature, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()

    # 3.2 工具层
    doc.add_heading('3.2 工具层（tools/）', level=2)

    doc.add_heading('工具基础设施', level=3)
    infra_features = [
        'Tool 抽象基类：约定 name / description / parameters / run()',
        'ToolResult 统一执行结果（ok / content / error / metadata）',
        'ToolRegistry 注册表：register / unregister / get / list_tools / call / acall',
        'get_schemas() 生成可直传 invoke_with_tools 的列表',
        'schema_hash() 提供集合指纹',
        '_safety：路径越界校验 / 危险命令黑名单',
    ]
    for feature in infra_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_heading('编程工具（14 个）', level=3)
    tools_table = doc.add_table(rows=15, cols=4, style='Light Grid Accent 1')
    tools_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tool_headers = ['工具类', 'name', '功能', '状态']
    for i, header in enumerate(tool_headers):
        tools_table.rows[0].cells[i].text = header
    tools_data = [
        ['ListFilesTool', 'list_files', '树形列目录，max_depth 截断，忽略 __pycache__ / .git / node_modules', '✓'],
        ['ReadFileTool', 'read_file', '带行号读取文件，start_line / end_line 切片，max_lines 截断', '✓'],
        ['SearchTextTool', 'search_text', '正则全文搜索，支持 case_sensitive / file_pattern glob 过滤', '✓'],
        ['WriteFileTool', 'write_file', '写入 UTF-8 文件，自动创建父目录，路径越界拒绝', '✓'],
        ['ApplyPatchTool', 'apply_patch', 'search-replace 替换：精确匹配 + 空白模糊匹配，原子写入', '✓'],
        ['MultiEditTool', 'multi_edit', '对单个文件按顺序应用多处 search-replace，全成功才落盘（原子）', '✓'],
        ['ReplaceLinesTool', 'replace_lines', '按 1 基闭区间 [start_line, end_line] 替换行，原子写入', '✓'],
        ['RunCommandTool', 'run_command', 'subprocess.run(shell=True)，含黑名单校验、timeout、stdout+stderr 合并', '✓'],
        ['WorkspaceInfoTool', 'workspace_info', '返回工作区绝对路径、顶层结构、主要语言、关键配置文件存在情况', '✓'],
        ['FindDefinitionTool', 'find_definition', '按符号名找定义，Universal Ctags 后端（覆盖 150+ 语言）', '✓'],
        ['SearchCodebaseTool', 'search_codebase', '自然语言/关键词/报错片段模糊搜代码，SQLite FTS5（零依赖）', '✓'],
        ['FindReferencesTool', 'find_references', '按符号名找引用，Tree-sitter 精确 AST，标注角色 def/call/import/attr', '✓'],
        ['RepoMapTool', 'repo_map', '生成「仓库地图」：各文件 top 符号清单，按符号数排名', '✓'],
        ['GetFileRelationsTool', 'get_file_relations', '查文件依赖：本文件 imports + 被谁 imported_by', '✓'],
    ]
    for i, row_data in enumerate(tools_data, 1):
        for j, cell_data in enumerate(row_data):
            tools_table.rows[i].cells[j].text = cell_data

    doc.add_heading('代码智能子系统 (tools/code_intel/)', level=3)
    code_intel_features = [
        '_common.py：语言路由 / 代码文件遍历 / 索引目录',
        'ctags.py：Universal Ctags 后端，CtagsIndex 生成/缓存 tags（mtime 过期重建）',
        'fts.py：SQLite FTS5 后端，FtsIndex 按行窗口分块入库、按 mtime 增量更新',
        'treesitter.py：Tree-sitter 后端，精确 AST 解析，带引用角色、import 依赖图',
    ]
    for feature in code_intel_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # 3.3 Agent 实现
    doc.add_heading('3.3 Agent 实现（agents/）', level=2)

    agent_capabilities = [
        ('CodingAgent 核心能力', [
            'ReAct 循环：调用 llm.invoke_with_tools → 有 tool_calls 则执行工具并把结果追回 → 无 tool_calls 即为最终答案',
            '多轮记忆：跨 run() 自动回灌历史',
            '同步执行：agent.run(task) 返回最终答案字符串',
            '异步执行：await agent.arun(task) 原生异步',
            '异步流式：async for ev in agent.arun_stream(task) 逐 AgentEvent 输出',
            '并发控制：max_concurrency（单步内工具并发上限，默认 4）',
            '超时控制：tool_timeout / llm_timeout（秒，None 不限）',
        ]),
        ('代码修改工作流', [
            '默认 system prompt 约束「读→计划→patch→验证→总结」五步',
            'require_confirm=True：所有改文件操作执行前请求确认',
            '确认时展示 diff/改动预览',
            'auto 模式：确认时输入 a 进入 auto 模式，后续不再询问；auto_approve=True 可直接启动',
        ]),
        ('健壮的代码编辑', [
            'apply_patch：精确 + 空白模糊匹配（唯一命中则替换并标注 fuzzy）',
            'multi_edit：一次多处替换，任一处失败则整体不写（原子）',
            'replace_lines：按行号区间替换，end_line 越界自动夹到文件末尾',
            '_edit_match.py：精确+空白模糊匹配、原子写入，供编辑类工具共享',
        ]),
        ('验证—修复闭环', [
            'verify_command 参数开启（如 "python -m unittest"）',
            'LLM 给出最终答案前、且本轮确实改过文件时，Agent 独立跑一遍验证',
            'parse_test_output 解析 unittest/pytest/compileall 输出（exit code 为准）',
            '失败则把结构化失败回灌为 user 消息继续修复',
            '最多 Config.max_fix_iters（默认 3）轮；达上限仍失败则接受答案并附提示',
            'trace 记 verify_run/verify_passed/verify_failed/verify_exhausted',
        ]),
        ('项目感知 + 改动追踪', [
            'project_context=True：detect_project_context() 检测语言/测试命令/关键文件/README',
            'format_project_context() 注入默认 system prompt',
            'track_edits=True：构造 EditTracker，在写文件工具前抓快照',
            '任务结束用 difflib 生成改动 diff',
            'get_last_diff_summary() 查询、并记入 trace edit_summary',
        ]),
        ('上下文预算（截断/压缩）', [
            '_prepare_messages() 调 compact_history 压缩超长跨轮历史',
            '_truncate_tool_output() 按 config.tool_output_max_* 截断喂回 LLM 的工具结果',
            '每步经 trace 记录 context_tokens',
            'enable_smart_compression 时用 LLM 摘要，否则 cheap 摘要',
        ]),
        ('稳健性（重试/熔断/预算）', [
            '所有 LLM 调用经 _call_llm_sync / _call_llm_async 包裹',
            '熔断检查（CircuitBreaker）+ 指数退避重试',
            'token_budget 设定单次任务 token 上限',
            '超限则 trace budget_exceeded 并跳出循环走兜底总结',
            '重试记 trace llm_retry',
            '默认 llm_max_retries=2、circuit_enabled=True、token_budget=None',
        ]),
        ('子 Agent 编排', [
            'enable_subagent=True 时注册 SubAgentTool',
            'LLM 可用 delegate_subtask(task) 把子任务交给独立子 Agent',
            '子 Agent 用局部导入避免循环依赖、enable_subagent=False 杜绝嵌套',
            '子 Agent 继承父级 require_confirm / verbose / auto_approve 安全姿态',
            '子 Agent 轻量模型支持：subagent_use_light_llm=True 时用轻量 LLM',
        ]),
        ('自动保存', [
            'config.auto_save_enabled 时重写的 add_message 每 auto_save_interval 条消息自动保存',
            '覆盖式保存 autosave-<name>.json（固定名覆盖、失败不影响主流程）',
            '记 trace auto_saved',
        ]),
        ('可选工具（P3）', [
            'todowrite_enabled → TodoWriteTool（todo_write，覆盖式待办清单，持久化 todo-<name>.json）',
            'devlog_enabled → DevLogTool（devlog，追加带时间戳的开发日志）',
            'skills_enabled → SkillTool（use_skill，列出/加载 skills_dir/*.md 技能）',
        ]),
        ('MCP 适配', [
            '构造参数 mcp_servers=[{name, command, args, env, cwd, timeout, prefix}]',
            '建 prompt 前为每个 spec 启动 MCPClient 并注册其工具',
            '运行期 add_mcp_server() 可热接',
            'close() 关闭所有 MCP 子进程（幂等）',
            'MCP 工具与内置工具同进 ToolRegistry，在 ReAct 循环里被 LLM 调用',
        ]),
        ('生产预设', [
            'CodingAgent.production() 类方法：一行拿到把全栈能力协同打开的 Agent',
            '默认开启：trace + 会话自动保存 + 项目自举 + 确认/auto + 改动追踪 + 子 Agent + 预算/并发',
            'Config 开关在 model_copy() 副本上设置（不污染传入的 config）',
            '**overrides 可覆盖',
        ]),
        ('文件读缓存', [
            'read_file 成功后经 _record_read 记录 {path: {size, mtime}}',
            '_get_read_cache() 覆盖基类返回真实读缓存并随会话落盘',
            'load_session 用 _detect_read_drift 比对磁盘 size/mtime，对漂移文件追加告警',
        ]),
    ]

    for cap_name, features in agent_capabilities:
        doc.add_heading(cap_name, level=3)
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # 第4章：测试覆盖情况
    doc.add_heading('4. 测试覆盖情况', level=1)

    doc.add_heading('测试概况', level=2)
    test_summary = [
        '测试框架: stdlib unittest（无需 pytest）',
        '测试用例: 246 个',
        '通过率: 100%',
        '外部依赖: 无（全部使用本地 StubLLM / 本地 fake MCP server）',
        '网络请求: 无（不发起任何真实 API 请求）',
    ]
    for item in test_summary:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('测试模块分布', level=2)
    test_table = doc.add_table(rows=14, cols=3, style='Light Grid Accent 1')
    test_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    test_headers = ['测试模块', '用例数', '覆盖范围']
    for i, header in enumerate(test_headers):
        test_table.rows[0].cells[i].text = header
    test_data = [
        ['test_tools.py', '18', '工具基础设施（Tool / ToolResult / ToolRegistry）'],
        ['test_coding_tools.py', '35', '9 个编程工具（含多处/行号/模糊编辑）'],
        ['test_coding_agent.py', '60+', 'CodingAgent 同步/异步/流式/会话/确认/auto 模式'],
        ['test_trace.py', '13', 'TraceLogger 单元 + CodingAgent 集成'],
        ['test_project_capabilities.py', '15', '项目上下文检测 + 改动追踪'],
        ['test_context_manager.py', '16', '上下文预算管理'],
        ['test_resilience.py', '13', '重试 / 熔断 / token 预算'],
        ['test_orchestration.py', '15', '子 Agent 委派 + 自动保存 + 轻量 LLM'],
        ['test_productivity_skills.py', '12', 'TodoWrite / DevLog / Skills'],
        ['test_production_preset.py', '3', 'CodingAgent.production() 生产预设'],
        ['test_mcp.py', '19', 'MCP 客户端 / 工具包装 / 集成'],
        ['test_verify.py', '11', '测试输出解析 + 验证闸门'],
        ['test_code_intel.py', '15', '代码智能五件套（ctags/tree-sitter/FTS5）'],
    ]
    for i, row_data in enumerate(test_data, 1):
        for j, cell_data in enumerate(row_data):
            test_table.rows[i].cells[j].text = cell_data

    doc.add_heading('测试特点', level=2)
    test_features = [
        '零外部依赖：所有测试使用本地 StubLLM，不发起真实 API 请求',
        '离线可跑：无需 .env 配置，无需网络连接',
        '覆盖全面：从工具单元测试到 Agent 集成测试，从同步到异步，从正常流程到异常处理',
        '代码智能测试按 ctags/tree-sitter 可用性自动 skip',
    ]
    for feature in test_features:
        doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # 第5章：Demo 示例
    doc.add_heading('5. Demo 示例', level=1)

    doc.add_heading('无需 .env 的 Demo（可直接运行）', level=2)
    demo_table1 = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
    demo_table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    demo_headers1 = ['Demo', '功能']
    for i, header in enumerate(demo_headers1):
        demo_table1.rows[0].cells[i].text = header
    demo_data1 = [
        ['coding_agent_resilience_demo.py', '稳健性演示（重试/熔断/预算，本地 FlakyLLM）'],
        ['coding_agent_orchestration_demo.py', '编排演示（子 Agent 委派 + 自动保存）'],
        ['coding_agent_productivity_demo.py', '生产力工具演示（todo/devlog/skills）'],
        ['coding_agent_mcp_demo.py', 'MCP 适配演示（第一段无需 .env）'],
    ]
    for i, row_data in enumerate(demo_data1, 1):
        for j, cell_data in enumerate(row_data):
            demo_table1.rows[i].cells[j].text = cell_data

    doc.add_heading('需要 .env 的 Demo（需配置 LLM API）', level=2)
    demo_table2 = doc.add_table(rows=9, cols=2, style='Light Grid Accent 1')
    demo_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    demo_headers2 = ['Demo', '功能']
    for i, header in enumerate(demo_headers2):
        demo_table2.rows[0].cells[i].text = header
    demo_data2 = [
        ['coding_agent_demo.py', '单次任务 demo'],
        ['coding_agent_chat_demo.py', '多轮对话 demo（逐步搭建项目）'],
        ['coding_agent_stream_demo.py', '异步流式 demo（逐事件实时输出）'],
        ['coding_agent_session_demo.py', '会话持久化 demo（保存/恢复/一致性检查）'],
        ['coding_agent_trace_demo.py', 'Trace demo（生成 JSONL + HTML 报告）'],
        ['coding_agent_advanced_demo.py', '高级能力 demo（项目感知 + 改动总结）'],
        ['coding_agent_context_demo.py', '上下文预算 demo（截断/压缩/token 计量）'],
        ['coding_agent_pro_demo.py', '生产预设 demo（一体化多轮项目编写）'],
    ]
    for i, row_data in enumerate(demo_data2, 1):
        for j, cell_data in enumerate(row_data):
            demo_table2.rows[i].cells[j].text = cell_data

    doc.add_page_break()

    # 第6章：验证记录
    doc.add_heading('6. 验证记录', level=1)

    doc.add_heading('测试执行结果', level=2)
    p = doc.add_paragraph()
    run = p.add_run(
        '$ python -m unittest discover -s tests\n'
        '----------------------------------------------------------------------\n'
        'Ran 246 tests in X.XXXs\n'
        '\n'
        'OK'
    )
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_heading('关键验证点', level=2)
    verification_points = [
        'Stage 0: 核心骨架，Agent 基类可实例化',
        'Stage 1-2: 工具基础设施 + 7 个编程工具',
        'Stage 3: CodingAgent ReAct 循环',
        'Stage 4: 代码修改闭环 + verbose/confirm',
        'Stage 5: 异步与流式输出',
        'Stage 6: 会话持久化与恢复',
        'Stage 7: 可观测性 Trace',
        'Stage 8: 增强编程能力（项目感知 + 改动总结）',
        'Stage 9: 上下文预算管理',
        'Stage 10: 稳健性与预算上限',
        'Stage 11: 子 Agent 编排 + 自动保存',
        'Stage 12: Skills / TodoWrite / DevLog',
        'P0-1: 代码编辑健壮性（apply_patch 模糊匹配 + multi_edit + replace_lines）',
        'P0-2: 跨语言符号检索/导航 → 升级为代码智能子系统 CI-1…CI-4',
        'P0-3: 验证—修复闭环',
        'MCP 适配: 作客户端消费外部 server',
    ]
    for point in verification_points:
        doc.add_paragraph(point, style='List Bullet')

    doc.add_page_break()

    # 第7章：项目亮点
    doc.add_heading('7. 项目亮点', level=1)

    highlights = [
        ('生产化就绪', [
            '完整的会话持久化与恢复机制，含环境一致性检查',
            '可观测性 Trace（JSONL + HTML 报告），敏感信息自动脱敏',
            '稳健性保障（重试/熔断/token 预算）',
            '一键生产预设 CodingAgent.production() 开箱即用',
        ]),
        ('代码智能能力', [
            '跨语言符号检索：Universal Ctags 后端覆盖 150+ 语言',
            '自然语言代码搜索：SQLite FTS5 零依赖实现',
            '精确 AST 解析：Tree-sitter 后端，带引用角色、import 依赖图',
            '仓库地图：全局符号概览，帮助 Agent 建立整体认知',
        ]),
        ('灵活的工具编辑', [
            '三种编辑模式：单处 apply_patch / 多处 multi_edit / 按行号 replace_lines',
            '精确匹配 + 空白模糊匹配，容错性强',
            '原子写入，任一处失败则整体不写',
            '确认时展示 diff/改动预览',
        ]),
        ('安全与可控', [
            '所有改文件操作 + 运行命令可配置确认',
            'auto 模式支持（输入 a 后续不再询问）',
            '子 Agent 继承父级安全姿态，避免「委派即绕过确认」',
            '路径越界校验 / 危险命令黑名单',
        ]),
        ('零外部依赖测试', [
            '246 个测试用例，100% 使用本地 StubLLM',
            '不发起任何真实 API 请求，不依赖外部网络',
            '代码智能测试按后端可用性自动 skip',
        ]),
        ('清晰的架构分层', [
            'core/ 与 tools/ 解耦，tools 仅依赖 core.exceptions',
            '工具层可独立先于 core 导入',
            'agents 依赖 core + tools，职责明确',
        ]),
    ]

    for highlight_name, features in highlights:
        doc.add_heading(highlight_name, level=2)
        for feature in features:
            doc.add_paragraph(feature, style='List Bullet')

    doc.add_page_break()

    # 第8章：后续规划
    doc.add_heading('8. 后续规划', level=1)

    doc.add_paragraph(
        '根据 docs/enhancement-roadmap.md 和 docs/coding-agent-roadmap.md，后续可扩展方向包括：'
    )

    future_plans = [
        '更多 LLM 后端：支持 Azure OpenAI、Anthropic Claude、本地模型等',
        '更丰富的工具：Git 操作、数据库查询、API 调用等',
        '更智能的上下文管理：基于重要性的消息筛选、自动摘要',
        '多 Agent 协作：Agent 间通信、任务分发、结果聚合',
        'MCP Server 角色：把自身工具暴露为 MCP server（当前仅作客户端）',
    ]
    for plan in future_plans:
        doc.add_paragraph(plan, style='List Bullet')

    doc.add_page_break()

    # 第9章：总结
    doc.add_heading('9. 总结', level=1)

    doc.add_paragraph(
        '本项目已完成路线图阶段 0–12 及 P0-1/P0-2/P0-3 的全部实现，交付了一个功能完整、'
        '架构清晰、测试充分的编程 Agent 核心库。246 个测试用例全部通过，具备生产化部署的基础能力。'
    )

    doc.add_heading('核心交付物', level=2)
    deliverables = [
        '14 个编程工具（含代码智能五件套）',
        '完整的 ReAct 循环 Agent（同步/异步/流式）',
        '会话持久化 + 可观测性 Trace + 上下文预算 + 稳健性保障',
        '子 Agent 编排 + MCP 适配 + 生产预设',
        '12 个可运行 Demo + 246 个测试用例',
    ]
    for item in deliverables:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('技术特色', level=2)
    tech_features = [
        '零外部依赖的代码智能实现（ctags + FTS5 + tree-sitter）',
        '原子化的代码编辑能力（精确 + 模糊匹配）',
        '完善的安全机制（确认/auto/路径校验/命令黑名单）',
        '生产级的可观测性与稳健性保障',
    ]
    for feature in tech_features:
        doc.add_paragraph(feature, style='List Bullet')

    # 保存文档
    output_path = r'E:\code\agent\learn_agent\docs\delivery-implementation-report.docx'
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")
    return output_path

if __name__ == '__main__':
    create_docx()